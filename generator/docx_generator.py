from docx import Document
from docx.shared import Inches
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls
from datetime import datetime
import tempfile
import os


def generate_order_docx(order_data, template_path="templates/primer.docx"):

    # Генерирует DOCX файл заказа из шаблона с правильными границами таблицы
    try:
        if not os.path.exists(template_path):
            raise FileNotFoundError(f"Шаблон не найден: {template_path}")

        doc = Document(template_path)

        # Заменяем плейсхолдеры в документе
        replace_placeholders(doc, order_data)

        # Создаем временный файл
        with tempfile.NamedTemporaryFile(delete=False, suffix='.docx') as tmp_file:
            output_path = tmp_file.name

        # Сохраняем документ
        doc.save(output_path)
        return output_path

    except Exception as e:
        raise Exception(f"Ошибка при генерации документа: {str(e)}")


def replace_placeholders(doc, order_data):

    # Заменяем простые плейсхолдеры
    for paragraph in doc.paragraphs:
        replace_in_paragraph(paragraph, order_data)

    # Обрабатываем таблицы
    for table in doc.tables:
        replace_in_table(table, order_data)


def replace_in_paragraph(paragraph, order_data):

    # Заменяет плейсхолдеры в одном параграфе
    for key, value in order_data.items():
        if key == 'items':
            continue

        placeholder = f"{{{{{key}}}}}"
        if placeholder in paragraph.text:
            paragraph.text = paragraph.text.replace(placeholder, str(value))


def replace_in_table(table, order_data):

    # Ищем строку с плейсхолдером товара
    for row_idx, row in enumerate(table.rows):
        for cell in row.cells:
            if "{{quantity}}" in cell.text or "{{product_id}}" in cell.text:
                # Нашли строку с товаром - заменяем её на все товары
                replace_items_in_table(table, row_idx, order_data['items'])
                return


def replace_items_in_table(table, template_row_idx, items):

    # Удаляем строку-шаблон
    table._tbl.remove(table.rows[template_row_idx]._tr)

    # Добавляем строки для каждого товара
    for item in items:
        new_row = table.add_row()

        # Устанавливаем границы для всей строки
        for cell in new_row.cells:
            # Устанавливаем границы для ячейки
            cell._element.tcPr.append(create_border_element())

        # Заполняем ячейки данными товара
        cells = new_row.cells
        if len(cells) >= 6:
            cells[0].text = str(item['quantity'])
            cells[1].text = str(item['product_id'])
            cells[2].text = item['product_name']
            cells[3].text = f"{item['wholesale_price']:.2f} руб."
            cells[4].text = f"{item['total_cost']:.2f} руб."
            cells[5].text = f"{item['retail_price']:.2f} руб."

            # Центрируем текст в ячейках
            for cell in cells:
                for paragraph in cell.paragraphs:
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER


def create_border_element():

    # Создает xml элемент с границами для ячейки таблицы

    border_xml = '''
    <w:tcPr xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">
        <w:tcBorders>
            <w:top w:val="single" w:sz="4" w:space="0" w:color="000000"/>
            <w:left w:val="single" w:sz="4" w:space="0" w:color="000000"/>
            <w:bottom w:val="single" w:sz="4" w:space="0" w:color="000000"/>
            <w:right w:val="single" w:sz="4" w:space="0" w:color="000000"/>
        </w:tcBorders>
    </w:tcPr>
    '''
    return parse_xml(border_xml)


def prepare_order_data(order, db_session):

    # Вычисляем общую стоимость всех товаров
    total_cost = 0
    items_data = []

    for item in order.items:
        item_total = item.product.price * item.quantity
        total_cost += item_total

        items_data.append({
            'quantity': item.quantity,
            'product_id': item.product.id_Product,
            'product_name': item.product.Name_tov,
            'wholesale_price': item.product.price,
            'total_cost': item_total,
            'retail_price': item.product.price * 1.2
        })

    order_data = {
        'order_code': order.Code,
        'buyer_name': order.buyer.Buyer_name if order.buyer else 'Не указан',
        'buyer_phone': order.buyer.number_buy if order.buyer else 'Не указан',
        'buyer_address': order.buyer.adress_buy if order.buyer else 'Не указан',
        'employee_name': order.employee.FIO if order.employee else 'Не указан',
        'order_date': datetime.now().strftime('%d.%m.%Y'),
        'order_status': order.status if order.status else 'В обработке',
        'order_total': f"{total_cost:.2f}",
        'description': order.description if order.description else 'Без дополнительных комментариев',
        'items': items_data
    }

    return order_data